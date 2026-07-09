# -*- coding: utf-8 -*-
"""
様式４ Excel出力モジュール

テンプレート構造（様式４（対策なし）シート）:
  印刷範囲: A1:AP50  用紙: A4縦  行高さ固定: 13.5pt（変更禁止）
  AP列(42) = 採点欄（点数1〜5）→ 書き込み禁止
  内容を書き込む列:
    工事名 G2  / 会社名 G3
    留意点: G列(7)  ← D列はラベル印刷済み（理由と頭文字位置を揃えるためG列開始）
    理由  : G列(7)  ← D:F列はラベル結合済み
  内容の右端: AO列(41) ← AP列は採点欄のため結合しない

  項目１ 行5-19:
    留意点① G5  / 理由①〜④ G6〜G9
    留意点② G10 / 理由①〜④ G11〜G14
    留意点③ G15 / 理由①〜④ G16〜G19
  項目２ 行20-34:
    留意点① G20 / 理由①〜④ G21〜G24
    留意点② G25 / 理由①〜④ G26〜G29
    留意点③ G30 / 理由①〜④ G31〜G34
  項目３ 行35-49:
    留意点① G35 / 理由①〜④ G36〜G39
    留意点② G40 / 理由①〜④ G41〜G44
    留意点③ G45 / 理由①〜④ G46〜G49
"""
import os, shutil, datetime, warnings
import openpyxl
from openpyxl.styles import Alignment, Font

# 全角カタカナ・記号 → 半角 変換テーブル
_KATA_HAN = {
    # ── カタカナ（清音・小文字）
    'ァ':'ｧ','ィ':'ｨ','ゥ':'ｩ','ェ':'ｪ','ォ':'ｫ',
    'ャ':'ｬ','ュ':'ｭ','ョ':'ｮ','ッ':'ｯ','ー':'ｰ',
    'ア':'ｱ','イ':'ｲ','ウ':'ｳ','エ':'ｴ','オ':'ｵ',
    'カ':'ｶ','キ':'ｷ','ク':'ｸ','ケ':'ｹ','コ':'ｺ',
    'サ':'ｻ','シ':'ｼ','ス':'ｽ','セ':'ｾ','ソ':'ｿ',
    'タ':'ﾀ','チ':'ﾁ','ツ':'ﾂ','テ':'ﾃ','ト':'ﾄ',
    'ナ':'ﾅ','ニ':'ﾆ','ヌ':'ﾇ','ネ':'ﾈ','ノ':'ﾉ',
    'ハ':'ﾊ','ヒ':'ﾋ','フ':'ﾌ','ヘ':'ﾍ','ホ':'ﾎ',
    'マ':'ﾏ','ミ':'ﾐ','ム':'ﾑ','メ':'ﾒ','モ':'ﾓ',
    'ヤ':'ﾔ','ユ':'ﾕ','ヨ':'ﾖ',
    'ラ':'ﾗ','リ':'ﾘ','ル':'ﾙ','レ':'ﾚ','ロ':'ﾛ',
    'ワ':'ﾜ','ヲ':'ｦ','ン':'ﾝ',
    # ── カタカナ（濁音・半濁音）
    'ガ':'ｶﾞ','ギ':'ｷﾞ','グ':'ｸﾞ','ゲ':'ｹﾞ','ゴ':'ｺﾞ',
    'ザ':'ｻﾞ','ジ':'ｼﾞ','ズ':'ｽﾞ','ゼ':'ｾﾞ','ゾ':'ｿﾞ',
    'ダ':'ﾀﾞ','ヂ':'ﾁﾞ','ヅ':'ﾂﾞ','デ':'ﾃﾞ','ド':'ﾄﾞ',
    'バ':'ﾊﾞ','ビ':'ﾋﾞ','ブ':'ﾌﾞ','ベ':'ﾍﾞ','ボ':'ﾎﾞ',
    'パ':'ﾊﾟ','ピ':'ﾋﾟ','プ':'ﾌﾟ','ペ':'ﾍﾟ','ポ':'ﾎﾟ',
    'ヴ':'ｳﾞ',
    # ── 全角記号（括弧類）
    '（':'(','）':')',
    '「':'｢','」':'｣',
    '『':'｢','』':'｣',
    '【':'[','】':']',
    '〔':'[','〕':']',
    # ── 全角記号（句読点・中点・長音）
    '・':'･','。':'｡','、':'､',
    # ── 全角記号（その他よく使う記号）
    '！':'!','？':'?',
    '／':'/','＼':'\\',
    '：':':','；':';',
    '～':'~','〜':'~',
    '＋':'+','－':'-','＊':'*','＝':'=',
    '＜':'<','＞':'>',
    '＆':'&','＠':'@','＃':'#','％':'%',
    '＿':'_','｜':'|',
    '…':'...','‥':'..',
}

def _zen2han(text: str) -> str:
    """全角カタカナ・記号を半角に変換する（ひらがな・漢字・英数字は変更しない）"""
    if not isinstance(text, str):
        return text
    return ''.join(_KATA_HAN.get(ch, ch) for ch in text)

_BASE = os.path.dirname(os.path.abspath(__file__))
SHEET = "様式４（対策なし）"

CELLS = {
    "工事名": (2, 7), "会社名": (3, 7),
    "項目１": [
        {"留意点": (5,  7), "理由": [(6,7),  (7,7),  (8,7),  (9,7)]},
        {"留意点": (10, 7), "理由": [(11,7), (12,7), (13,7), (14,7)]},
        {"留意点": (15, 7), "理由": [(16,7), (17,7), (18,7), (19,7)]},
    ],
    "項目２": [
        {"留意点": (20, 7), "理由": [(21,7), (22,7), (23,7), (24,7)]},
        {"留意点": (25, 7), "理由": [(26,7), (27,7), (28,7), (29,7)]},
        {"留意点": (30, 7), "理由": [(31,7), (32,7), (33,7), (34,7)]},
    ],
    "項目３": [
        {"留意点": (35, 7), "理由": [(36,7), (37,7), (38,7), (39,7)]},
        {"留意点": (40, 7), "理由": [(41,7), (42,7), (43,7), (44,7)]},
        {"留意点": (45, 7), "理由": [(46,7), (47,7), (48,7), (49,7)]},
    ],
}
ITEM_KEYS = ["項目１", "項目２", "項目３"]


def export(project_name: str, company_name: str, proposal: dict, out_dir: str = None) -> str:
    # テンプレートパスを毎回解決
    _t1 = os.path.join(_BASE, "template", "様式４テンプレート.xlsx")
    _t2 = os.path.join(_BASE, "様式４テンプレート.xlsx")
    template = _t1 if os.path.exists(_t1) else (_t2 if os.path.exists(_t2) else None)
    if template is None:
        raise FileNotFoundError(
            "Excelテンプレートファイルが見つかりません。\n"
            "「様式４テンプレート.xlsx」をアプリフォルダまたはtemplateフォルダに配置してください。"
        )

    out_dir = out_dir or os.path.join(_BASE, "data")
    os.makedirs(out_dir, exist_ok=True)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = "".join(c for c in project_name if c not in r'\/:*?"<>|')[:25]
    path = os.path.join(out_dir, f"様式４_{safe}_{ts}.xlsx")

    shutil.copy2(template, path)
    warnings.filterwarnings("ignore")
    wb = openpyxl.load_workbook(path)
    if SHEET not in wb.sheetnames:
        raise ValueError(
            f"テンプレートにシート「{SHEET}」が見つかりません。\n"
            f"存在するシート: {', '.join(wb.sheetnames)}\n"
            f"テンプレートファイルを確認してください。"
        )
    ws = wb[SHEET]

    # 内容右端: AO列(41)  ← AP列(42)は採点欄のため絶対に結合しない
    _END = 41  # AO

    # 留意点行を G:AO 結合（理由と頭文字位置を揃えるためG列開始）
    for row in [5, 10, 15, 20, 25, 30, 35, 40, 45]:
        ws.merge_cells(start_row=row, start_column=7, end_row=row, end_column=_END)

    # 理由行を G:AO 結合（D:F ラベル結合はテンプレート既存のため触らない）
    for row in (list(range(6,  10)) + list(range(11, 15)) + list(range(16, 20)) +
                list(range(21, 25)) + list(range(26, 30)) + list(range(31, 35)) +
                list(range(36, 40)) + list(range(41, 45)) + list(range(46, 50))):
        ws.merge_cells(start_row=row, start_column=7, end_row=row, end_column=_END)

    # フォント: 留意点・工事名行=UD明朝、理由行=UDP明朝
    _FONT_CHUI = "BIZ UD明朝 Medium"
    _FONT_RIYU = "BIZ UDP明朝 Medium"
    _CHUI_ROWS = {2, 3, 5, 10, 15, 20, 25, 30, 35, 40, 45}

    def w(row, col, val):
        cell = ws.cell(row=row, column=col)
        cell.value = _zen2han(val)
        # フォント名のみ設定し、サイズはテンプレートの値を絶対に変更しない
        font_name = _FONT_CHUI if row in _CHUI_ROWS else _FONT_RIYU
        cell.font = cell.font.copy(name=font_name)
        # 行高さはテンプレート固定値(13.5pt)を維持するため変更しない
        # shrink_to_fitは絶対に使わない（文字サイズを変えてはいけないため）
        cell.alignment = Alignment(wrap_text=False, vertical="center", shrink_to_fit=False)

    r, c = CELLS["工事名"]; w(r, c, project_name)
    r, c = CELLS["会社名"]; w(r, c, company_name)

    for i, key in enumerate(list(proposal.keys())[:3]):
        entries = proposal.get(key, [])
        for j, cell_info in enumerate(CELLS[ITEM_KEYS[i]]):
            entry = entries[j] if j < len(entries) else {}
            w(*cell_info["留意点"], entry.get("留意点", ""))
            for k, (rr, rc) in enumerate(cell_info["理由"]):
                reasons = entry.get("理由リスト", [])
                w(rr, rc, reasons[k] if k < len(reasons) else "")

    wb.save(path)
    wb.close()
    return path


def _set_docx_cell(cell, text, font_name):
    """docx セルへフォント指定で書き込む（東アジアフォントも設定）。ラベル・体裁は保持し内容のみ差し込む。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = cell.paragraphs[0]
    # 既存ランを全て除去（テンプレートの空ラン含む）してから内容を差し込む
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(_zen2han(text) if text else "")
    run.font.name = font_name
    run.font.size = Pt(10.5)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)


def export_word(project_name: str, company_name: str, proposal: dict, out_dir: str = None) -> str:
    from docx import Document

    # Wordテンプレートパスを毎回解決（Excelと同様 template/ 配下優先）
    _t1 = os.path.join(_BASE, "template", "様式４テンプレート.docx")
    _t2 = os.path.join(_BASE, "様式４テンプレート.docx")
    template = _t1 if os.path.exists(_t1) else (_t2 if os.path.exists(_t2) else None)
    if template is None:
        raise FileNotFoundError(
            "Wordテンプレートファイルが見つかりません。\n"
            "「様式４テンプレート.docx」をアプリフォルダまたはtemplateフォルダに配置してください。"
        )

    out_dir = out_dir or os.path.join(_BASE, "data")
    os.makedirs(out_dir, exist_ok=True)
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = "".join(c for c in project_name if c not in r'\/:*?"<>|')[:25]
    path = os.path.join(out_dir, f"様式４_{safe}_{ts}.docx")

    _FONT_NOTE = "BIZ UD明朝 Medium"    # 留意点・工事名・会社名
    _FONT_RIYU = "BIZ UDP明朝 Medium"   # 理由

    doc = Document(template)
    if len(doc.tables) < 2:
        raise ValueError("Wordテンプレートの表構成が不正です（情報表・本体表の2表が必要）。")

    info = doc.tables[0]   # 0行目=工事名, 1行目=会社名（列1が値セル）
    body = doc.tables[1]   # 45行 × 3列（列2が内容セル）

    # 工事名・会社名
    _set_docx_cell(info.cell(0, 1), project_name, _FONT_NOTE)
    _set_docx_cell(info.cell(1, 1), company_name, _FONT_NOTE)

    # 項目1〜3 × 留意点①②③（各: 留意点1行 + 理由4行）
    item_keys = list(proposal.keys())[:3]
    for i, key in enumerate(item_keys):
        base = i * 15
        entries = proposal.get(key, [])
        for g in range(3):                       # 留意点①②③
            note_row = base + g * 5
            entry = entries[g] if g < len(entries) else {}
            _set_docx_cell(body.cell(note_row, 2), entry.get("留意点", ""), _FONT_NOTE)
            reasons = entry.get("理由リスト", [])
            for k in range(4):                   # 理由①〜④
                r_text = reasons[k] if k < len(reasons) else ""
                _set_docx_cell(body.cell(note_row + 1 + k, 2), r_text, _FONT_RIYU)

    doc.save(path)
    return path
